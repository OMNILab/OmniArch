"""
File Processing Utilities
Contains functions for processing different file formats and extracting text content
"""

import streamlit as st
import os
import tempfile
from typing import Optional, Dict, Any
import re


def extract_text_from_file(uploaded_file) -> Optional[str]:
    """
    Extract text content from uploaded file based on its format.

    Args:
        uploaded_file: Streamlit uploaded file object

    Returns:
        str: Extracted text content or None if extraction fails
    """
    if not uploaded_file:
        return None

    file_name = uploaded_file.name.lower()
    file_extension = os.path.splitext(file_name)[1].lower()

    try:
        if file_extension == ".txt":
            return _extract_text_from_txt(uploaded_file)
        elif file_extension == ".md":
            return _extract_text_from_markdown(uploaded_file)
        elif file_extension == ".docx":
            return _extract_text_from_docx(uploaded_file)
        elif file_extension == ".pdf":
            return _extract_text_from_pdf(uploaded_file)
        else:
            st.error(f"不支持的文件格式: {file_extension}")
            return None

    except Exception as e:
        st.error(f"文件处理失败: {str(e)}")
        return None


def _extract_text_from_txt(uploaded_file) -> str:
    """Extract text from TXT file"""
    try:
        # Try UTF-8 first
        content = uploaded_file.read().decode("utf-8")
        uploaded_file.seek(0)  # Reset file pointer
        return content
    except UnicodeDecodeError:
        try:
            # Try GBK for Chinese text
            uploaded_file.seek(0)
            content = uploaded_file.read().decode("gbk")
            uploaded_file.seek(0)
            return content
        except UnicodeDecodeError:
            # Try with error handling
            uploaded_file.seek(0)
            content = uploaded_file.read().decode("utf-8", errors="ignore")
            uploaded_file.seek(0)
            return content


def _extract_text_from_markdown(uploaded_file) -> str:
    """Extract text from Markdown file"""
    # Markdown files are essentially text files
    content = _extract_text_from_txt(uploaded_file)

    # Clean up markdown formatting for better processing
    content = _clean_markdown_content(content)

    return content


def _extract_text_from_docx(uploaded_file) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document

        # Save uploaded file to temporary location
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_file_path = tmp_file.name

        try:
            # Read the document
            doc = Document(tmp_file_path)

            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)

        finally:
            # Clean up temporary file
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)

    except ImportError:
        st.error("需要安装 python-docx 库来处理 DOCX 文件")
        return None
    except Exception as e:
        st.error(f"DOCX 文件处理失败: {str(e)}")
        return None


def _extract_text_from_pdf(uploaded_file) -> str:
    """Extract text from PDF file with better Chinese character support"""
    try:
        # Save uploaded file to temporary location
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_file_path = tmp_file.name

        try:
            # Try multiple PDF libraries in order of preference
            text_parts = []

            # Method 1: Try pdfplumber first (best for Chinese)
            try:
                import pdfplumber

                with pdfplumber.open(tmp_file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(page_text.strip())

                if text_parts:
                    content = "\n\n".join(text_parts)
                    # Check if we got meaningful Chinese text
                    chinese_chars = [
                        char for char in content if "\u4e00" <= char <= "\u9fff"
                    ]
                    if len(chinese_chars) > 5:  # If we have reasonable Chinese content
                        return content
            except Exception as e:
                pass  # Fall back to next method

            # Method 2: Try pymupdf (fitz)
            try:
                import fitz

                doc = fitz.open(tmp_file_path)
                text_parts = []

                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())

                doc.close()

                if text_parts:
                    content = "\n\n".join(text_parts)
                    # Check if we got meaningful Chinese text
                    chinese_chars = [
                        char for char in content if "\u4e00" <= char <= "\u9fff"
                    ]
                    if len(chinese_chars) > 5:  # If we have reasonable Chinese content
                        return content
            except Exception as e:
                pass  # Fall back to next method

            # Method 3: Try PyPDF2 as fallback
            try:
                import PyPDF2

                text_parts = []

                with open(tmp_file_path, "rb") as file:
                    pdf_reader = PyPDF2.PdfReader(file)

                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(page_text.strip())

                if text_parts:
                    content = "\n\n".join(text_parts)
                    # Even if Chinese chars are garbled, return the content
                    # User can see there's an encoding issue
                    return content
            except Exception as e:
                pass

            # If all methods fail, return empty string
            return ""

        finally:
            # Clean up temporary file
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)

    except ImportError:
        st.error("需要安装 PDF 处理库。请运行: pip install pdfplumber pymupdf PyPDF2")
        return None
    except Exception as e:
        st.error(f"PDF 文件处理失败: {str(e)}")
        return None


def _clean_markdown_content(content: str) -> str:
    """Clean markdown formatting for better text processing"""
    if not content:
        return content

    # Remove markdown headers
    content = re.sub(r"^#{1,6}\s+", "", content, flags=re.MULTILINE)

    # Remove markdown links but keep text
    content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", content)

    # Remove markdown bold/italic formatting
    content = re.sub(r"\*\*([^*]+)\*\*", r"\1", content)
    content = re.sub(r"\*([^*]+)\*", r"\1", content)
    content = re.sub(r"__([^_]+)__", r"\1", content)
    content = re.sub(r"_([^_]+)_", r"\1", content)

    # Remove markdown code blocks
    content = re.sub(r"```[^`]*```", "", content, flags=re.DOTALL)
    content = re.sub(r"`([^`]+)`", r"\1", content)

    # Remove markdown lists
    content = re.sub(r"^[\s]*[-*+]\s+", "", content, flags=re.MULTILINE)
    content = re.sub(r"^[\s]*\d+\.\s+", "", content, flags=re.MULTILINE)

    # Remove horizontal rules
    content = re.sub(r"^[\s]*[-*_]{3,}[\s]*$", "", content, flags=re.MULTILINE)

    # Clean up extra whitespace
    content = re.sub(r"\n\s*\n", "\n\n", content)
    content = re.sub(r"[ \t]+", " ", content)

    return content.strip()


def get_supported_file_types() -> Dict[str, str]:
    """Get supported file types and their descriptions"""
    return {
        "txt": "纯文本文件 (.txt)",
        "md": "Markdown 文档 (.md)",
        "docx": "Word 文档 (.docx)",
        "pdf": "PDF 文档 (.pdf)",
    }


def validate_file_size(uploaded_file, max_size_mb: int = 10) -> bool:
    """
    Validate file size

    Args:
        uploaded_file: Streamlit uploaded file object
        max_size_mb: Maximum file size in MB

    Returns:
        bool: True if file size is valid, False otherwise
    """
    if not uploaded_file:
        return False

    # Get file size in bytes
    uploaded_file.seek(0, 2)  # Seek to end
    file_size = uploaded_file.tell()
    uploaded_file.seek(0)  # Reset to beginning

    # Convert to MB
    file_size_mb = file_size / (1024 * 1024)

    if file_size_mb > max_size_mb:
        st.error(f"文件大小超过限制 ({file_size_mb:.1f}MB > {max_size_mb}MB)")
        return False

    return True


def get_file_info(uploaded_file) -> Dict[str, Any]:
    """
    Get information about the uploaded file

    Args:
        uploaded_file: Streamlit uploaded file object

    Returns:
        dict: File information
    """
    if not uploaded_file:
        return {}

    file_name = uploaded_file.name
    file_extension = os.path.splitext(file_name)[1].lower()

    # Get file size
    uploaded_file.seek(0, 2)
    file_size = uploaded_file.tell()
    uploaded_file.seek(0)

    file_size_mb = file_size / (1024 * 1024)

    return {
        "name": file_name,
        "extension": file_extension,
        "size_bytes": file_size,
        "size_mb": file_size_mb,
        "type": get_supported_file_types().get(file_extension[1:], "未知格式"),
    }
